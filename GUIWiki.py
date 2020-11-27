# GUIWiki.py
import wikipedia
#python to doc
from docx import Document
def Wiki(keyword,lang='th'):
                wikipedia.set_lang(lang)
                # summary = บทความที่สรุป
                data = wikipedia.summary(keyword)
                # page + content = บทความทั้งหน้า
                data2 = wikipedia.page(keyword)
                data2 = data2.content
                doc = Document() # สร้างไฟล์ word ด้วย python
                doc.add_heading(keyword,0)
                doc.add_paragraph(data2)
                doc.save(keyword+'.docx')
                print('สร้างไฟล์สำเร็จ')

#เปลี่ยนเป็นภาษาไทย
wikipedia.set_lang('th')
from tkinter import *
from tkinter import ttk
from tkinter import messagebox



GUI = Tk()

GUI.title('โปรแกรม WIKI')
GUI.geometry('400x300')
v_search = StringVar()
#Config
FONT1 = ('Angsana New',15)
#คำอธิบาย
L1 = ttk.Label(GUI,text = 'ค้นหาบทความ',font=FONT1)
L1.pack()

#ช่องกรอกข้อมูล
E1 = ttk.Entry(GUI,textvariable = v_search,font=FONT1,width=40)
E1.pack(pady=10)

def Search():
                keyword = v_search.get() #ดึงข้อมูลเข้ามา
                try:
                                #ลองค้นหาดูว่าได้ผลลัพธ์หรือไม่ หากได้ให้ผ่าน
                                language = v_radio.get()
                                Wiki(keyword,language)
                                messagebox.showinfo('บันทึกสำเร็จ','ค้นหาและบันทึกข้อมูลสำเร็จแล้ว')
                except:
                                messagebox.showwarning('key word error','กรุณากรอกคำค้นหาใหม่')
                                
                #print(wikipedia.search(keyword))
                #result = wikipedia.summary(keyword)
                #print(result)
                
#ปุ่มค้นหา
B1 = ttk.Button(GUI,text='Search',command=Search)
B1.pack(ipadx=20,ipady=10)

# เลือกภาษา
F1 = Frame(GUI)
F1.pack(pady=50)

v_radio = StringVar() # ช่องเก็บข้อมูลภาษา

RB1 = ttk.Radiobutton(F1,text = 'ภาษาไทย',variable=v_radio,value='th')
RB2 = ttk.Radiobutton(F1,text = 'ภาษาอังกฤษ',variable=v_radio,value='en')
RB3 = ttk.Radiobutton(F1,text = 'ภาษาจีน',variable=v_radio,value='zh')
RB1.invoke()#สั่งให้ค่าเริ่มต้นเป็นภาษาไทย

RB1.grid(row = 0,column = 0)
RB2.grid(row = 0,column = 1)
RB3.grid(row = 0,column = 2)

GUI.mainloop()
